#!/usr/bin/env python3
"""
expand_zvit2.py — Розширення Звіт_практики_фінальний2.docx:
  - Тиждень 2: формальна модель планування + схема даних Redis
  - Тиждень 3: рис 3.12/3.13 (.py і git workflow), секція 2.3 обробка помилок,
               секція 4.5 cold/warm бенчмарк (рис 3.14)
  - Додаток Б: лістинг ключових фрагментів коду
  - Додаток В: повний перелік тестів pytest
  - Додаток Г: UML-діаграми (плейсхолдери)
"""

import os
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

DOCX_IN  = '/home/mika/projects/gpu-job-scheduler/Звіт_практики_фінальний2.docx'
DOCX_OUT = '/home/mika/projects/gpu-job-scheduler/Звіт_практики_фінальний3.docx'

doc = Document(DOCX_IN)

# ─────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────

def find_para(doc, fragment):
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None

def add_para(text='', bold=False, italic=False, center=False, style='Normal'):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p._p

def add_placeholder(description):
    """Сірий курсивний плейсхолдер для нового рисунку."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ СКРІНШОТ: {description} ]')
    run.italic = True
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '888888')
    rPr.append(color_el)
    return p._p

def chain_after(ref_el, elements):
    """Вставляє список _p елементів по черзі після ref_el."""
    cur = ref_el
    for el in elements:
        cur.addnext(el)
        cur = el

def insert_before(ref_para, el):
    ref_para._p.addprevious(el)

# ─────────────────────────────────────────────────────────────────────
# ТИЖДЕНЬ 2 — ФОРМАЛЬНА МОДЕЛЬ ПЛАНУВАННЯ
# Вставляємо перед "Стани задачі та переходи..."
# ─────────────────────────────────────────────────────────────────────

states_para = find_para(doc, 'Стани задачі та переходи між ними наведено')
if states_para:
    model_paras = [
        add_para('Формальна модель та алгоритм вибору GPU', bold=True),
        add_para(),
        add_para(
            'Формально систему планування визначено наступним чином. '
            'Нехай G = {g₀, g₁, ..., gₙ₋₁} — множина доступних GPU, '
            'де n = GPU_COUNT. Кожна GPU gᵢ у будь-який момент часу '
            'перебуває в одному зі станів: idle або busy. '
            'Черга задач Q реалізована як Redis-список job_queue, '
            'де кожен елемент — UUID задачі у порядку FIFO.'
        ),
        add_para(),
        add_para(
            'Алгоритм вибору GPU для задачі j (вираз 5):'
        ),
        add_para(
            'for i in range(GPU_COUNT):       # перебір GPU за індексом   (5)'
        ),
        add_para(
            '    with r.pipeline() as pipe:'
        ),
        add_para(
            '        pipe.watch(f"gpu:{i}")'
        ),
        add_para(
            '        gpu = json.loads(pipe.get(f"gpu:{i}"))'
        ),
        add_para(
            '        if gpu["status"] == "idle":'
        ),
        add_para(
            '            pipe.multi()'
        ),
        add_para(
            '            gpu["status"] = "busy"'
        ),
        add_para(
            '            pipe.set(f"gpu:{i}", json.dumps(gpu))'
        ),
        add_para(
            '            pipe.execute()  # → WatchError або успіх'
        ),
        add_para(
            '            return i         # повертаємо індекс GPU'
        ),
        add_para(
            'де pipe — Redis Pipeline з підтримкою WATCH; '
            'WatchError означає конкурентну модифікацію ключа; '
            'при WatchError цикл повторюється з початку.'
        ),
        add_para(),
        add_para(
            'Обрана стратегія називається First-Idle (перша вільна). '
            'Вона гарантує детерміновано-справедливий розподіл: '
            'задачі виконуються в порядку надходження (FIFO-черга), '
            'а GPU призначаються за мінімальним вільним індексом. '
            'При нульовій довжині черги середній час очікування '
            'W = 0; при довжині черги k і n GPU із середнім часом '
            'виконання T: W ≈ ⌈k/n⌉ · T. '
            'Для кафедрального сервера з двома GPU і типовим k ≤ 4 '
            'очікування не перевищує одного циклу виконання.'
        ),
        add_para(),
        add_para('Схема зберігання даних у Redis', bold=True),
        add_para(),
        add_para(
            'Redis виконує роль як брокера повідомлень (job_queue), '
            'так і сховища стану. Всього використовується п\'ять типів ключів (табл. 13).'
        ),
        add_para(),
        add_para('Таблиця 13 — Структура ключів Redis', bold=False),
        add_para('Ключ | Тип | Формат / Значення | TTL'),
        add_para('job:{uuid} | String | JSON: job_id, status, student_id, runtime, cpus, memory, gpu_id, code, stdout, stderr, artifacts, queued_at, started_at, finished_at | Без TTL'),
        add_para('gpu:{0..n-1} | String | JSON: gpu_id, status (idle|busy), job_id, updated_at | Без TTL'),
        add_para('job:logs:{uuid} | List | Рядки виводу: "stdout:текст" або "stderr:текст" | Без TTL'),
        add_para('job_queue | List | Список UUID задач у порядку FIFO (RPUSH / BLPOP) | Без TTL'),
        add_para('apikey:{key} | String | student_id — ідентифікатор студента, якому належить ключ | Може мати expires_at'),
        add_para(),
        add_para(
            'Відсутність TTL для job:* є свідомим рішенням: задачі і їх логи '
            'зберігаються до явного виклику DELETE /jobs/cleanup, '
            'що дає студентам час завантажити артефакти. '
            'Ключі gpu:* ініціалізуються Worker при старті з усіма статусами idle '
            'і оновлюються виключно через атомарні транзакції.'
        ),
        add_para(),
    ]
    chain_after(states_para._p.getprevious(), model_paras)
    print('  [NEW] Тиждень 2: формальна модель + Redis schema')

# ─────────────────────────────────────────────────────────────────────
# ТИЖДЕНЬ 3 — РИС. 3.12 (подача .py файлу) після опису другого способу
# ─────────────────────────────────────────────────────────────────────

py_para = find_para(doc, 'Другий спосіб (.py + .zip) призначений')
if py_para:
    fig312 = [
        add_para(),
        add_placeholder(
            'Веб-форма з вкладкою "Файл": вибрано файл train.py, '
            'завантажено dataset.zip, вибрано pytorch-cu121 — '
            'кнопка Submit активна; поруч — сторінка статусу з логом '
            '"== installing requirements ==" та "== running user code =="'
        ),
        add_para(
            'Рисунок 3.12 — Подача задачі через завантаження .py файлу та .zip датасету',
            italic=True, center=True
        ),
        add_para(),
    ]
    chain_after(py_para._p, fig312)
    print('  [NEW] Рис. 3.12 placeholder (.py файл)')

# ─────────────────────────────────────────────────────────────────────
# ТИЖДЕНЬ 3 — РИС. 3.13 (git-репозиторій) після опису третього способу
# ─────────────────────────────────────────────────────────────────────

git_para = find_para(doc, 'Третій спосіб (git-репозиторій) підходить')
if git_para:
    fig313 = [
        add_para(),
        add_placeholder(
            'Веб-форма: поле "Git repo URL" заповнено посиланням на GitHub, '
            'поле entrypoint = "train.py"; сторінка статусу показує '
            'лог "Cloning into /workspace/...", потім вивід навчання'
        ),
        add_para(
            'Рисунок 3.13 — Подача задачі через git-репозиторій (URL + entrypoint)',
            italic=True, center=True
        ),
        add_para(),
    ]
    chain_after(git_para._p, fig313)
    print('  [NEW] Рис. 3.13 placeholder (git repo)')

# ─────────────────────────────────────────────────────────────────────
# ТИЖДЕНЬ 3 — СЕКЦІЯ 2.3: ОБРОБКА ПОМИЛОК ТА ВІДНОВЛЕННЯ СТАНУ
# Вставляємо після секції 2.2 (після пояснення _stream_reader)
# ─────────────────────────────────────────────────────────────────────

stream_para = find_para(doc, 'де stream — об\'єкт stdout або stderr процесу')
if stream_para:
    err_section = [
        add_para(),
        add_para('2.3. Обробка помилок та відновлення стану Worker', bold=True),
        add_para(),
        add_para(
            'Виконання довільного студентського коду вимагає надійної '
            'обробки будь-яких збоїв: помилок у коді, перевищення таймауту, '
            'запиту на скасування та аварійного перезапуску Worker. '
            'Реалізовано три механізми захисту.'
        ),
        add_para(),
        add_para('Таймаут виконання', bold=True),
        add_para(
            'Функція execute_job містить моніторинговий цикл, що щосекунди '
            'перевіряє тривалість виконання контейнера. Якщо вона перевищила '
            'JOB_TIMEOUT_SECONDS (за замовчуванням 300 с), Worker надсилає '
            'docker kill і записує у Redis статус failed із поясненням '
            '"Job timed out after Ns". Такий підхід захищає від нескінченних '
            'циклів та задач, що зависли в очікуванні введення.'
        ),
        add_para(),
        add_para('Скасування задачі користувачем', bold=True),
        add_para(
            'При виклику POST /jobs/{id}/cancel Master встановлює у Redis '
            'прапор cancel_requested: true. Моніторинговий цикл Worker '
            'читає цей прапор і при його наявності надсилає docker kill — '
            'незалежно від того, скільки часу задача вже виконується. '
            'Стан переходить у cancelled, GPU звільняється.'
        ),
        add_para(),
        add_para('Блок finally та очищення ресурсів', bold=True),
        add_para(
            'Незалежно від результату виконання (успіх, помилка, '
            'скасування, таймаут), блок finally гарантує: '
            '1) звільнення GPU через release_gpu(gpu_id); '
            '2) видалення тимчасових файлів з диска хоста '
            '(код .py, клонований репозиторій, requirements.txt). '
            'Такий підхід виключає витік ресурсів при будь-якому '
            'сценарії завершення задачі.'
        ),
        add_para(),
        add_para('Відновлення після перезапуску Worker', bold=True),
        add_para(
            'При кожному старті Worker виконує функцію recover_stale_gpus(). '
            'Вона перевіряє всі GPU зі статусом busy у Redis і для кожного '
            'виконує docker inspect: якщо відповідний контейнер не запущений, '
            'GPU повертається у стан idle, а задача отримує статус failed '
            'із повідомленням "Worker restarted — job was lost". '
            'Це вирішує проблему зависших резервацій після аварійного '
            'зупинення або оновлення Worker.'
        ),
        add_para(),
        add_para(
            'Загальна логіка обробки помилок у Worker відповідає '
            'принципу fail-fast: будь-яке необроблене виключення '
            'перехоплюється у зовнішньому try/except execute_job, '
            'записується у Redis як статус failed, і GPU звільняється. '
            'Worker продовжує обслуговувати чергу незалежно від '
            'помилок в окремих задачах.'
        ),
        add_para(),
    ]
    chain_after(stream_para._p, err_section)
    print('  [NEW] Тиждень 3: секція 2.3 обробка помилок')

# ─────────────────────────────────────────────────────────────────────
# ТИЖДЕНЬ 3 — СЕКЦІЯ 4.5: БЕНЧМАРК ХОЛОДНОГО ТА ТЕПЛОГО СТАРТУ
# Вставляємо після "Рисунок 3.5 — Результати автоматизованого аудиту"
# ─────────────────────────────────────────────────────────────────────

audit_fig = find_para(doc, 'Рисунок 3.5 — Результати автоматизованого аудиту')
if audit_fig:
    cold_section = [
        add_para(),
        add_para('4.5. Бенчмарк холодного та теплого старту контейнера', bold=True),
        add_para(),
        add_para(
            'Окрім lifecycle-бенчмарку загального циклу задачі, '
            'проведено спеціалізоване вимірювання часу старту контейнера: '
            'перший запуск (cold start) та наступні (warm start). '
            'Відмінність між ними обумовлена кешуванням Docker: '
            'після першого запуску шари образу залишаються в кеші '
            'файлової системи хоста, і наступні запуски використовують '
            'кешовані дані без повторної розпаковки.'
        ),
        add_para(),
        add_para(
            'Бенчмарк cold_warm.py виконав 1 cold та 5 warm запусків '
            'runtime-профілю pytorch-cu121 з мінімальним тестовим скриптом '
            '(print GPU info + torch.cuda.is_available()). '
            'Метрика startup_s — час від HTTP POST /submit-job до '
            'переходу задачі у статус running. '
            'Результати наведено в табл. 14 та на рис. 3.14.'
        ),
        add_para(),
        add_para('Таблиця 14 — Результати бенчмарку cold/warm start (pytorch-cu121)', bold=False),
        add_para('Запуск | Тип | Час черги (s) | Загальний час (s)'),
        add_para('1 | cold | 0.009 | 2.016'),
        add_para('2 | warm | 0.015 | 2.022'),
        add_para('3 | warm | 0.012 | 2.019'),
        add_para('4 | warm | 0.011 | 2.520'),
        add_para('5 | warm | 0.009 | 2.015'),
        add_para('6 | warm | 0.009 | 5.243'),
        add_para(),
        add_para(
            'Час черги (startup_s) стабільно становить 9–15 мс незалежно '
            'від типу запуску — це підтверджує, що Docker-образ вже '
            'завантажений і затримка зумовлена лише Redis round-trip. '
            'У реальному розгортанні перший запуск нового образу '
            'потребує docker pull (~500 МБ для pytorch-cu121), '
            'що може зайняти 1–5 хвилин залежно від каналу. '
            'Тому на кафедральному сервері рекомендується попередньо '
            'завантажити всі образи командою docker compose pull.'
        ),
        add_para(),
        add_placeholder(
            'cold_warm.png: стовпчастий графік, де по осі X — номери запусків (1–6), '
            'перший стовпець позначено "cold" (рожевий), решта "warm" (синій), '
            'по осі Y — загальний час у секундах'
        ),
        add_para(
            'Рисунок 3.14 — Бенчмарк холодного та теплого старту контейнера',
            italic=True, center=True
        ),
        add_para(),
    ]
    chain_after(audit_fig._p, cold_section)
    print('  [NEW] Тиждень 3: секція 4.5 cold/warm benchmark')

# ─────────────────────────────────────────────────────────────────────
# ДОДАТОК Б — ЛІСТИНГ КЛЮЧОВИХ ФРАГМЕНТІВ КОДУ
# Вставляємо після останнього абзацу Додатку А
# ─────────────────────────────────────────────────────────────────────

last_admin_para = find_para(doc, 'GET /admin (HTMLResponse)')
if last_admin_para:
    appendix_b = [
        add_para(),
        add_para('Додаток Б', style='Heading 1'),
        add_para('ЛІСТИНГ КЛЮЧОВИХ ФРАГМЕНТІВ КОДУ', bold=True),
        add_para(),
        add_para(
            'У цьому додатку наведено повний лістинг ключових функцій '
            'системи: AST-аналізатор безпеки, SSE-стрімінг логів, '
            'атомарна резервація GPU, стрімер потоків контейнера, '
            'головна функція виконання задачі та основний цикл Worker.'
        ),
        add_para(),

        # ── Б.1 check_code_safety ──
        add_para('Б.1. Функція AST-аналізу безпеки коду (master.py)', bold=True),
        add_para(),
        add_para('DANGEROUS_BUILTINS = {"exec", "eval", "compile", "__import__"}'),
        add_para(),
        add_para('def check_code_safety(code: str) -> tuple[bool, str]:'),
        add_para('    if len(code.encode()) > 200 * 1024:'),
        add_para('        return False, "Code exceeds 200 KB size limit"'),
        add_para('    try:'),
        add_para('        tree = ast.parse(code)'),
        add_para('    except SyntaxError as e:'),
        add_para('        return False, f"Syntax error: {e}"'),
        add_para('    for node in ast.walk(tree):'),
        add_para('        if isinstance(node, ast.Call) and isinstance(node.func, ast.Name):'),
        add_para('            if node.func.id in DANGEROUS_BUILTINS:'),
        add_para('                return False, f"Use of {node.func.id}() is not allowed"'),
        add_para('    return True, ""'),
        add_para(),
        add_para(
            'Функція отримує рядок коду, перевіряє розмір (не більше 200 КБ), '
            'парсить його через ast.parse() і обходить синтаксичне дерево. '
            'При виявленні виклику будь-якої функції зі списку DANGEROUS_BUILTINS '
            'повертає False з описом порушення. '
            'SyntaxError у коді студента також відхиляється — '
            'це запобігає подачі некоректного коду в чергу.'
        ),
        add_para(),

        # ── Б.2 SSE endpoint ──
        add_para('Б.2. SSE-ендпоінт стрімінгу логів (master.py)', bold=True),
        add_para(),
        add_para('@app.get("/jobs/{job_id}/logs/stream")'),
        add_para('async def stream_logs(job_id: str, request: Request):'),
        add_para('    raw = await asyncio.to_thread(r.get, f"job:{job_id}")'),
        add_para('    if not raw:'),
        add_para('        raise HTTPException(status_code=404, detail="Job not found")'),
        add_para(),
        add_para('    async def event_generator():'),
        add_para('        sent = 0'),
        add_para('        while True:'),
        add_para('            if await request.is_disconnected():'),
        add_para('                break'),
        add_para('            new_entries = await asyncio.to_thread('),
        add_para('                r.lrange, f"job:logs:{job_id}", sent, -1)'),
        add_para('            for entry in new_entries:'),
        add_para('                yield {"data": entry}'),
        add_para('                sent += 1'),
        add_para('            job_raw = await asyncio.to_thread(r.get, f"job:{job_id}")'),
        add_para('            if job_raw:'),
        add_para('                job = json.loads(job_raw)'),
        add_para('                if job["status"] in ("completed","failed","cancelled"):'),
        add_para('                    tail = await asyncio.to_thread('),
        add_para('                        r.lrange, f"job:logs:{job_id}", sent, -1)'),
        add_para('                    for entry in tail:'),
        add_para('                        yield {"data": entry}'),
        add_para('                    yield {"data": json.dumps({"type": "end",'),
        add_para('                                               "status": job["status"]})}'),
        add_para('                    return'),
        add_para('            await asyncio.sleep(0.4)'),
        add_para(),
        add_para('    return EventSourceResponse(event_generator())'),
        add_para(),
        add_para(
            'Ендпоінт використовує asyncio та EventSourceResponse (бібліотека sse-starlette). '
            'Генератор опитує Redis-список job:logs:{id} кожні 0.4 с, '
            'надсилає нові рядки клієнту у форматі SSE і завершується після отримання '
            'фінального статусу задачі. Перевірка is_disconnected() '
            'запобігає накопиченню "мертвих" з\'єднань.'
        ),
        add_para(),

        # ── Б.3 get_free_gpu / reserve ──
        add_para('Б.3. Атомарна резервація GPU (worker.py)', bold=True),
        add_para(),
        add_para('def get_free_gpu() -> int | None:'),
        add_para('    for gpu_id in range(GPU_COUNT):'),
        add_para('        key = f"gpu:{gpu_id}"'),
        add_para('        try:'),
        add_para('            with r.pipeline() as pipe:'),
        add_para('                pipe.watch(key)'),
        add_para('                gpu = json.loads(pipe.get(key) or "{}")'),
        add_para('                if gpu.get("status") != "idle":'),
        add_para('                    pipe.reset()'),
        add_para('                    continue'),
        add_para('                pipe.multi()'),
        add_para('                gpu["status"]     = "busy"'),
        add_para('                gpu["updated_at"] = datetime.now(timezone.utc).isoformat()'),
        add_para('                pipe.set(key, json.dumps(gpu))'),
        add_para('                pipe.execute()'),
        add_para('                return gpu_id'),
        add_para('        except WatchError:'),
        add_para('            continue'),
        add_para('    return None'),
        add_para(),
        add_para(
            'WATCH встановлює спостереження за ключем gpu:{id}. '
            'Якщо між WATCH і EXECUTE інший потік змінив ключ, '
            'Redis скасовує транзакцію і кидає WatchError. '
            'Цикл продовжується далі до наступного GPU. '
            'Функція повертає індекс першої вільної GPU '
            'або None якщо всі зайняті.'
        ),
        add_para(),

        # ── Б.4 _stream_reader ──
        add_para('Б.4. Функція читання потоків контейнера (worker.py)', bold=True),
        add_para(),
        add_para('def _stream_reader(stream, stream_type: str,'),
        add_para('                   lines_list: list, log_key: str):'),
        add_para('    """Thread target: reads subprocess stream line-by-line."""'),
        add_para('    for raw_line in stream:'),
        add_para('        line = raw_line.rstrip("\\n")'),
        add_para('        lines_list.append(line)'),
        add_para('        r.rpush(log_key, f"{stream_type}:{line}")'),
        add_para(),
        add_para(
            'Запускається у двох потоках паралельно для stdout і stderr. '
            'Паралельне читання запобігає deadlock: якщо читати '
            'stdout і stderr послідовно в одному потоці, '
            'заповнення буфера одного потоку заблокує запис іншого. '
            'Кожен рядок одразу публікується у Redis-список '
            'з префіксом "stdout:" або "stderr:".'
        ),
        add_para(),

        # ── Б.5 recover_stale_gpus ──
        add_para('Б.5. Відновлення зависших резервацій GPU (worker.py)', bold=True),
        add_para(),
        add_para('def recover_stale_gpus():'),
        add_para('    for gpu_id in range(GPU_COUNT):'),
        add_para('        key = f"gpu:{gpu_id}"'),
        add_para('        gpu = json.loads(r.get(key) or "{}")'),
        add_para('        if gpu.get("status") != "busy":'),
        add_para('            continue'),
        add_para('        job_id = gpu.get("job_id")'),
        add_para('        container_name = f"job-{job_id}" if job_id else None'),
        add_para('        container_running = False'),
        add_para('        if container_name:'),
        add_para('            result = subprocess.run('),
        add_para('                ["docker", "inspect", "--format",'),
        add_para('                 "{{.State.Running}}", container_name],'),
        add_para('                capture_output=True, text=True)'),
        add_para('            container_running = result.stdout.strip() == "true"'),
        add_para('        if not container_running:'),
        add_para('            release_gpu(gpu_id)'),
        add_para('            if job_id:'),
        add_para('                job = get_job(job_id)'),
        add_para('                if job and job.get("status") == "running":'),
        add_para('                    update_job(job_id, {'),
        add_para('                        "status": "failed",'),
        add_para('                        "stderr": "Worker restarted — job was lost",'),
        add_para('                        "finished_at": datetime.now(timezone.utc).isoformat(),'),
        add_para('                    })'),
        add_para(),
        add_para(
            'Функція виконується одноразово при кожному старті Worker. '
            'Для кожної GPU зі статусом busy перевіряється наявність '
            'відповідного контейнера через docker inspect. '
            'Якщо контейнер не знайдено або не запущено — GPU звільняється, '
            'задача маркується як failed. '
            'Це вирішує проблему неконсистентного стану після аварійного '
            'зупинення або Docker daemon restart.'
        ),
        add_para(),

        # ── Б.6 main() loop ──
        add_para('Б.6. Головний цикл Worker (worker.py)', bold=True),
        add_para(),
        add_para('def main():'),
        add_para('    check_real_gpu_count()'),
        add_para('    recover_stale_gpus()'),
        add_para('    while True:'),
        add_para('        item = r.blpop("job_queue", timeout=5)'),
        add_para('        if not item:'),
        add_para('            continue'),
        add_para('        _, job_id = item'),
        add_para('        job = json.loads(r.get(f"job:{job_id}") or "{}")'),
        add_para('        if job.get("status") == "cancelled":'),
        add_para('            continue   # задачу скасовано до виконання'),
        add_para('        gpu_id = None'),
        add_para('        while gpu_id is None:'),
        add_para('            if get_job(job_id).get("cancel_requested"):'),
        add_para('                break'),
        add_para('            gpu_id = get_free_gpu()'),
        add_para('            if gpu_id is None:'),
        add_para('                time.sleep(1)'),
        add_para('        if gpu_id is None:'),
        add_para('            continue'),
        add_para('        t = threading.Thread('),
        add_para('            target=execute_job, args=(job_id, gpu_id), daemon=True)'),
        add_para('        t.start()'),
        add_para(),
        add_para(
            'BLPOP блокує потік до появи елемента у черзі з таймаутом 5 с. '
            'Таймаут потрібен щоб перевіряти стан черги і не блокуватись назавжди. '
            'Після отримання задачі Worker чекає вільну GPU у циклі polling, '
            'перевіряючи cancel_requested. '
            'Кожна задача запускається у daemon-потоці threading.Thread, '
            'що дозволяє паралельно обслуговувати кілька GPU. '
            'Daemon-потоки завершуються разом з основним процесом при зупиненні Worker.'
        ),
        add_para(),
    ]
    chain_after(last_admin_para._p, appendix_b)
    print('  [NEW] Додаток Б: лістинг коду')

# ─────────────────────────────────────────────────────────────────────
# ДОДАТОК В — ПОВНИЙ ПЕРЕЛІК ТЕСТІВ PYTEST
# ─────────────────────────────────────────────────────────────────────

TESTS = [
    ('TestHealth', [
        ('test_health_returns_ok', 'GET /health повертає status=ok'),
        ('test_runtimes_list', 'GET /runtimes повертає всі три профілі'),
        ('test_gpus_list', 'GET /gpus повертає список GPU зі статусами'),
        ('test_cluster_status_has_new_fields', 'GET /cluster-status містить queue_length та jobs'),
    ]),
    ('TestJobSubmit', [
        ('test_submit_valid_job', 'POST /submit-job з коректними параметрами → 201'),
        ('test_submit_with_name_and_description', 'Поля name та description зберігаються у Redis'),
        ('test_submit_with_student_id', 'student_id у тілі запиту приймається'),
        ('test_submit_invalid_runtime', 'Невідомий runtime → 400'),
        ('test_submit_empty_code', 'Порожній код → 400'),
        ('test_submit_invalid_cpus', 'cpus > 8.0 → 422'),
        ('test_submit_invalid_memory', 'memory без суфікса → 400'),
    ]),
    ('TestCodeSafety', [
        ('test_eval_blocked', 'eval("...") у коді → 400'),
        ('test_exec_blocked', 'exec("...") у коді → 400'),
        ('test_compile_blocked', 'compile() у коді → 400'),
        ('test_dunder_import_blocked', '__import__("os") у коді → 400'),
        ('test_syntax_error_rejected', 'синтаксична помилка у коді → 400'),
    ]),
    ('TestRequirementsSafety', [
        ('test_simple_requirements_accepted', 'numpy==1.26 у requirements → прийнято'),
        ('test_editable_install_blocked', '-e ./local у requirements → 400'),
        ('test_index_url_blocked', '--index-url у requirements → 400'),
        ('test_git_install_blocked', 'git+ URL у requirements → 400'),
        ('test_find_links_blocked', '--find-links у requirements → 400'),
    ]),
    ('TestSubmitJobForm', [
        ('test_form_with_pasted_code', 'POST /submit-job-form з кодом у полі code → 201'),
        ('test_form_with_py_file', 'завантаження .py файлу → задача з файлом'),
        ('test_form_non_py_file_rejected', 'файл не .py → 400'),
        ('test_form_no_code_no_file_no_repo', 'порожня форма → 400'),
        ('test_form_invalid_repo_url', 'некоректний URL репо → 400'),
        ('test_form_github_url_accepted', 'github.com URL → прийнято'),
        ('test_form_with_zip_dataset', '.zip датасет разом із .py → завантажено'),
        ('test_form_dataset_not_zip_rejected', 'датасет не .zip → 400'),
    ]),
    ('TestQueueLimit', [
        ('test_queue_limit_returns_429', 'MAX_QUEUED_PER_STUDENT задач → 429 на N+1'),
        ('test_running_job_counts_toward_limit', 'задача running враховується в ліміті'),
    ]),
    ('TestJobExecution', [
        ('test_job_completes_successfully', 'простий print() → status=completed'),
        ('test_job_with_error_fails', 'raise RuntimeError → status=failed'),
        ('test_gpu_is_available_in_container', 'torch.cuda.is_available() = True в контейнері'),
        ('test_job_status_transitions', 'queued → running → completed послідовно'),
    ]),
    ('TestArtifacts', [
        ('test_artifact_saved_and_listed', 'файл у OUTPUT_DIR → з\'являється у /artifacts'),
        ('test_artifact_download', 'GET /artifacts/file.csv → завантажується'),
        ('test_artifact_not_found', 'GET /artifacts/missing → 404'),
        ('test_artifacts_endpoint_returns_empty_list', 'без артефактів → порожній список'),
    ]),
    ('TestDataset', [
        ('test_data_dir_env_var_set_when_dataset_uploaded', 'DATA_DIR доступний у контейнері'),
    ]),
    ('TestRequirementsInstallation', [
        ('test_pip_install_runs_before_code', '"== installing requirements ==" у логах'),
    ]),
    ('TestLogStream', [
        ('test_logs_endpoint_for_completed_job', 'GET /logs/stream повертає рядки + done'),
        ('test_logs_endpoint_unknown_job', 'GET /logs/stream для невідомої задачі → 404'),
    ]),
    ('TestJobCancellation', [
        ('test_cancel_long_running_job', 'POST /cancel під час running → cancelled'),
        ('test_cancel_finished_job', 'POST /cancel після completed → 400'),
    ]),
    ('TestMonitoring', [
        ('test_job_appears_in_list', 'задача з\'являється у GET /jobs'),
        ('test_jobs_filter_by_status', 'фільтрація за статусом працює'),
        ('test_get_job_not_found', 'GET /jobs/unknown → 404'),
        ('test_queue_position_field_for_queued', 'поле position у відповіді для queued'),
    ]),
    ('TestCleanup', [
        ('test_cleanup_removes_finished_jobs', 'DELETE /jobs/cleanup видаляє completed/failed'),
    ]),
    ('TestAdminAuth', [
        ('test_admin_check_with_valid_key', 'GET /admin/check з правильним ключем → 200'),
        ('test_admin_check_with_invalid_key', 'невірний ключ → 401'),
        ('test_admin_check_no_key', 'без ключа → 401'),
        ('test_admin_stats_structure', 'GET /admin/stats містить by_student та by_runtime'),
    ]),
    ('TestApiKeys', [
        ('test_create_list_and_revoke_key', 'POST /admin/keys → GET /admin/keys → DELETE'),
        ('test_create_key_requires_student_id', 'POST без student_id → 422'),
        ('test_keys_endpoints_require_admin_auth', 'без admin key → 401'),
        ('test_submit_with_api_key_overrides_form_student_id', 'API-ключ перекриває student_id форми'),
    ]),
    ('TestAdminKill', [
        ('test_admin_kill_long_running_job', 'POST /admin/jobs/{id}/kill → status=failed'),
        ('test_admin_kill_unknown_job', 'kill невідомої задачі → 404'),
        ('test_admin_kill_without_auth_rejected', 'kill без admin key → 401'),
    ]),
    ('TestUI', [
        ('test_dashboard_renders', 'GET / повертає 200 та HTML'),
        ('test_admin_page_renders', 'GET /admin повертає 200 та HTML'),
    ]),
]

total = sum(len(t) for _, t in TESTS)

last_code_para = find_para(doc, 'Daemon-потоки завершуються разом з основним процесом')
if not last_code_para:
    last_code_para = find_para(doc, 'Б.6. Головний цикл Worker')

if last_code_para:
    appendix_v = [
        add_para(),
        add_para('Додаток В', style='Heading 1'),
        add_para(f'ПОВНИЙ ПЕРЕЛІК АВТОМАТИЗОВАНИХ ТЕСТІВ PYTEST ({total} ТЕСТІВ)', bold=True),
        add_para(),
        add_para(
            f'У цьому додатку наведено повний перелік {total} автоматизованих тестів '
            'у 16 класах. Усі тести розташовані у файлі tests/test_api.py '
            'і виконуються командою pytest tests/ -v. '
            'Середній час виконання повного набору — близько 3–4 хвилин '
            '(тести TestJobExecution та TestArtifacts запускають реальні Docker-контейнери).'
        ),
        add_para(),
    ]
    for class_name, tests in TESTS:
        appendix_v.append(add_para(class_name, bold=True))
        for func_name, desc in tests:
            appendix_v.append(add_para(f'    {func_name} — {desc}'))
        appendix_v.append(add_para())

    chain_after(last_code_para._p, appendix_v)
    print(f'  [NEW] Додаток В: {total} тестів')

# ─────────────────────────────────────────────────────────────────────
# ДОДАТОК Г — UML-ДІАГРАМИ
# ─────────────────────────────────────────────────────────────────────

last_test_para = find_para(doc, 'test_admin_page_renders')
if not last_test_para:
    last_test_para = find_para(doc, 'TestUI')

if last_test_para:
    appendix_g = [
        add_para(),
        add_para('Додаток Г', style='Heading 1'),
        add_para('UML-ДІАГРАМИ ТА СХЕМА ДАНИХ', bold=True),
        add_para(),

        # ── Г.1 Sequence diagram ──
        add_para('Г.1. Sequence-діаграма повного запиту', bold=True),
        add_para(),
        add_para(
            'На рисунку Г.1 наведено послідовність взаємодії компонентів '
            'від моменту надсилання студентом задачі до отримання результату. '
            'Учасники діаграми: Browser (клієнт), Master (FastAPI), '
            'Redis (черга та стан), Worker (виконавець), Docker (контейнер).'
        ),
        add_para(),
        add_para('Кроки sequence-діаграми:'),
        add_para('1. Browser → Master: POST /submit-job-form (code, runtime, cpus, memory)'),
        add_para('2. Master → Master: check_code_safety(code) — AST-аналіз'),
        add_para('3. Master → Redis: SET job:{uuid} {status:queued, ...}'),
        add_para('4. Master → Redis: RPUSH job_queue {uuid}'),
        add_para('5. Master → Browser: 201 {job_id, status:queued, position}'),
        add_para('6. Browser → Master: GET /jobs/{id}/logs/stream (SSE)'),
        add_para('7. Worker → Redis: BLPOP job_queue (блокуючий pop)'),
        add_para('8. Worker → Redis: WATCH/MULTI/EXEC gpu:{i} (резервація GPU)'),
        add_para('9. Worker → Redis: SET job:{id} {status:running, gpu_id:i}'),
        add_para('10. Worker → Docker: docker run --gpus device=i ... python -u script.py'),
        add_para('11. Docker → Worker: stdout/stderr рядки (readline loop)'),
        add_para('12. Worker → Redis: RPUSH job:logs:{id} "stdout:рядок"'),
        add_para('13. Master → Redis: LRANGE job:logs:{id} sent -1 (polling 0.4s)'),
        add_para('14. Master → Browser: SSE data: рядок (реальний час)'),
        add_para('15. Docker → Worker: exit code 0'),
        add_para('16. Worker → Redis: SET job:{id} {status:completed, artifacts:[...]}'),
        add_para('17. Worker → Redis: SET gpu:{i} {status:idle}'),
        add_para('18. Master → Browser: SSE data: {"type":"end","status":"completed"}'),
        add_para(),
        add_placeholder(
            'Sequence-діаграма UML: вертикальні лінії Browser / Master / Redis / Worker / Docker, '
            'горизонтальні стрілки між ними відповідно до 18 кроків, '
            'кроки 11-14 у циклі "loop [while running]"'
        ),
        add_para('Рисунок Г.1 — Sequence-діаграма повного циклу виконання задачі',
                 italic=True, center=True),
        add_para(),

        # ── Г.2 Class diagram ──
        add_para('Г.2. Діаграма класів Pydantic-моделей (master.py)', bold=True),
        add_para(),
        add_para(
            'На рисунку Г.2 наведено class-діаграму Pydantic-моделей, '
            'які визначають структуру вхідних даних REST API. '
            'Валідація виконується автоматично FastAPI при десеріалізації запиту.'
        ),
        add_para(),
        add_para('Клас JobRequest (POST /submit-job):'),
        add_para('    code: str — Python-код (обов\'язкове)'),
        add_para('    student_id: str = "anonymous"'),
        add_para('    runtime: str = "pytorch-cu121"'),
        add_para('    cpus: float = 2.0  (0.1 ≤ cpus ≤ 8.0)'),
        add_para('    memory: str = "2g"  (regex: \\d+(m|g|mb|gb))'),
        add_para('    name: str = ""'),
        add_para('    description: str = ""'),
        add_para('    requirements: str = ""'),
        add_para(),
        add_para('Клас ApiKeyCreate (POST /admin/keys):'),
        add_para('    student_id: str'),
        add_para('    expires_days: int = 365'),
        add_para(),
        add_para('Внутрішня структура job-запису у Redis (dict):'),
        add_para('    job_id, status, student_id, runtime, image'),
        add_para('    cpus, memory, gpu_id, code, repo_url'),
        add_para('    has_dataset, entrypoint, requirements'),
        add_para('    queued_at, started_at, finished_at'),
        add_para('    stdout, stderr, return_code, artifacts'),
        add_para('    cancel_requested, container_name'),
        add_para(),
        add_placeholder(
            'Class-діаграма UML: прямокутники JobRequest, ApiKeyCreate з атрибутами та типами; '
            'окремий блок "Redis Job Record" з усіма полями; '
            'зв\'язок JobRequest → Redis Job Record (create)'
        ),
        add_para('Рисунок Г.2 — Діаграма класів Pydantic-моделей та структура Redis-запису',
                 italic=True, center=True),
        add_para(),

        # ── Г.3 Redis key space ──
        add_para('Г.3. Повна схема ключового простору Redis', bold=True),
        add_para(),
        add_para(
            'На рисунку Г.3 наведено entity-діаграму всіх типів ключів '
            'Redis із зазначенням типу структури даних, '
            'мультиплікатора (кількість ключів) та зв\'язків між ними.'
        ),
        add_para(),
        add_para('job:{uuid}  [String/JSON]  — по одному на кожну задачу'),
        add_para('    ↑ посилається на → gpu:{i}  [String/JSON]  — n ключів (n = GPU_COUNT)'),
        add_para('job:logs:{uuid}  [List]  — по одному на задачу, зв\'язаний з job:{uuid}'),
        add_para('job_queue  [List]  — один глобальний список UUID'),
        add_para('apikey:{key}  [String]  — по одному на кожен API-ключ студента'),
        add_para(),
        add_para(
            'Зв\'язок між job:{uuid} і gpu:{i} реалізований через поле '
            'job_id у JSON gpu:{i} (Worker записує туди UUID задачі '
            'при резервації). Зворотний зв\'язок — через поле gpu_id '
            'у job:{uuid} (записується при переході в running). '
            'Ця двостороння прив\'язка дозволяє як Worker знайти задачу '
            'по GPU, так і клієнту отримати GPU задачі по job_id.'
        ),
        add_para(),
        add_placeholder(
            'Entity-діаграма: п\'ять прямокутників з назвами ключів, '
            'стрілки між job:{uuid} → gpu:{i} та job:{uuid} → job:logs:{uuid}, '
            'стрілка job_queue → job:{uuid} (dequeue), '
            'типи Redis зазначені в дужках під кожним прямокутником'
        ),
        add_para('Рисунок Г.3 — Схема ключового простору Redis',
                 italic=True, center=True),
        add_para(),
    ]
    chain_after(last_test_para._p, appendix_g)
    print('  [NEW] Додаток Г: UML-діаграми')

# ─────────────────────────────────────────────────────────────────────
# ЗБЕРЕЖЕННЯ
# ─────────────────────────────────────────────────────────────────────

doc.save(DOCX_OUT)
print(f'\nГотово! Збережено: {DOCX_OUT}')
