#!/usr/bin/env python3
"""
expand_zvit.py — Розширення Звіт_практики_фінальний.docx:
  - вставка PNG-графіків перед підписами рисунків
  - плейсхолдери для скріншотів
  - повний текст ВИСНОВКИ
  - новий підрозділ 1.3 веб-інтерфейс
  - розширення розділу 1.2
  - Додаток А з документацією API
  - розширення тижня 4
"""

import os
import re
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

DOCX_IN  = '/home/mika/projects/gpu-job-scheduler/Звіт_практики_фінальний.docx'
PLOTS    = '/home/mika/projects/gpu-job-scheduler/benchmarks/results/plots/'
DOCX_OUT = '/home/mika/projects/gpu-job-scheduler/Звіт_практики_фінальний_v2.docx'

doc = Document(DOCX_IN)

# ─────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────

def find_para(doc, fragment):
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None

def add_para(text='', bold=False, italic=False, center=False):
    """Append paragraph to document end and return its _p element."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p._p

def add_placeholder(description):
    """Gray italic placeholder. Returns _p."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ СКРІНШОТ: {description} ]')
    run.italic = True
    # gray color via XML
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '888888')
    rPr.append(color_el)
    return p._p

def add_image_para(image_path, width_inches=5.5):
    """Paragraph with centered image. Returns _p."""
    tmp = Document()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(width_inches))
    p_copy = deepcopy(p._p)
    return p_copy

def chain_insert_after(ref_el, elements):
    """
    Insert list of _p elements after ref_el in document order (top→bottom).
    Uses chaining so order is preserved.
    """
    cur = ref_el
    for el in elements:
        cur.addnext(el)
        cur = el

def insert_before(ref_p, el):
    ref_p._p.addprevious(el)

# ─────────────────────────────────────────────────────────────────────
# 1. ВСТАВКА ІСНУЮЧИХ PNG ПЕРЕД ПІДПИСАМИ РИСУНКІВ
# ─────────────────────────────────────────────────────────────────────

png_map = {
    'Рисунок 2.1': 'system_architecture.png',
    'Рисунок 2.3': 'job_lifecycle.png',
    'Рисунок 3.1': 'sse_vs_polling.png',
    'Рисунок 3.2': 'runtime_compare.png',
    'Рисунок 3.3': 'lifecycle.png',
    'Рисунок 3.4': 'concurrency.png',
    'Рисунок 3.5': 'security_audit.png',
}

for caption_frag, fname in png_map.items():
    para = find_para(doc, caption_frag)
    if para:
        img_el = add_image_para(os.path.join(PLOTS, fname), width_inches=5.3)
        insert_before(para, img_el)
        print(f'  [IMG] {caption_frag} ← {fname}')
    else:
        print(f'  [WARN] не знайдено: {caption_frag}')

# ─────────────────────────────────────────────────────────────────────
# 2. ПЛЕЙСХОЛДЕРИ ДЛЯ РИСУНКІВ БЕЗ PNG
# ─────────────────────────────────────────────────────────────────────

placeholder_map = {
    'Рисунок 1.1': 'Порівняння VM / Docker / LXC — таблиця або блок-схема з трьома стовпцями за критеріями ізоляції, накладних витрат, доступу до GPU',
    'Рисунок 1.2': 'Архітектура NVIDIA Container Toolkit: шари nvidia-container-runtime → libnvidia-container → GPU-пристрої хоста',
    'Рисунок 1.3': 'Конфлікти версій CUDA: до контейнеризації (перетинання символічних посилань) vs після (ізольовані стеки у кожному контейнері)',
    'Рисунок 2.2': 'Docker Compose: сервіси master, worker, redis — їх порти, мережа та точки монтування томів',
    'Рисунок 2.4': 'Атомарна резервація GPU: послідовність WATCH → GET → MULTI → SET → EXECUTE, гілки WatchError → retry та success → GPU busy',
    'Рисунок 2.5': 'Defense in Depth: 7 концентричних рубежів від зовнішнього (AST-аналіз) до внутрішнього (API-ключі)',
}

for caption_frag, desc in placeholder_map.items():
    para = find_para(doc, caption_frag)
    if para:
        ph_el = add_placeholder(desc)
        insert_before(para, ph_el)
        print(f'  [PH]  {caption_frag}')
    else:
        print(f'  [WARN] не знайдено: {caption_frag}')

# ─────────────────────────────────────────────────────────────────────
# 3. ПРИБИРАЄМО "(вставити ...)" ІЗ ТЕКСТУ
# ─────────────────────────────────────────────────────────────────────

for para in doc.paragraphs:
    if '(вставити ' in para.text:
        for run in para.runs:
            if '(вставити ' in run.text:
                run.text = re.sub(r'\s*\(вставити[^)]+\)', '', run.text)
        print(f'  [FIX] прибрано "(вставити...)" з: {para.text[:60]}...')

# ─────────────────────────────────────────────────────────────────────
# 4. ВИПРАВЛЕННЯ ОБІРВАНОГО АБЗАЦУ "Параметр --gpus device=N"
# ─────────────────────────────────────────────────────────────────────

stray = find_para(doc, 'Параметр --gpus device=N')
if stray:
    # Замінити на повноцінний текст
    for run in stray.runs:
        run.text = ''
    stray.add_run(
        'Параметр --gpus device=N є ключовим елементом ізоляції GPU: '
        'він обмежує видимість GPU для контейнера рівно однією фізичною '
        'картою, що унеможливлює конфлікти між паралельними задачами різних '
        'студентів. Повний формат команди запуску контейнера наведено у '
        'виразі 4.'
    )
    print('  [FIX] виправлено обірваний абзац "Параметр --gpus device=N"')

# ─────────────────────────────────────────────────────────────────────
# 5. РОЗШИРЕННЯ РОЗДІЛУ 1.2 «ТРИ СПОСОБИ ПОДАЧІ ЗАДАЧІ»
# ─────────────────────────────────────────────────────────────────────

git_para = find_para(doc, 'Git-репозиторій — Worker виконує')
if git_para:
    expansion = [
        add_para(),
        add_para(
            'Незалежно від способу подачі, задача проходить однаковий '
            'pipeline: AST-аналіз коду → постановка в Redis-чергу → '
            'атомарна резервація GPU → запуск Docker-контейнера → '
            'SSE-стрімінг логів → збір артефактів. Результати зберігаються '
            'у директорії /workspace/output контейнера і доступні через '
            'GET /jobs/{id}/artifacts після завершення.'
        ),
        add_para(),
        add_para(
            'Перший спосіб (Monaco Editor / JSON) оптимальний для коротких '
            'скриптів і швидкого прототипування: код вставляється в '
            'браузерний редактор Monaco з підсвіткою Python і відправляється '
            'до POST /submit-job як JSON. Ліміт розміру — 200 КБ.'
        ),
        add_para(),
        add_para(
            'Другий спосіб (.py + .zip) призначений для більших проєктів. '
            '.zip-датасет перевіряється на zip-slip атаку та розпаковується '
            'у /workspace/data/. Якщо разом із .py файлом завантажено '
            'requirements.txt, Worker перед виконанням запускає pip install '
            'у режимі --network bridge (єдиний виняток із правила мережевої '
            'ізоляції).'
        ),
        add_para(),
        add_para(
            'Третій спосіб (git-репозиторій) підходить для командних '
            'проєктів і відтворюваних ML-експериментів. Worker виконує '
            'git clone --depth 1 для мінімізації часу і диску, після чого '
            'запускає вказаний entrypoint-файл. Підтримуються публічні '
            'посилання GitHub та GitLab.'
        ),
        add_para(),
    ]
    chain_insert_after(git_para._p, expansion)
    print('  [EXP] розширено розділ 1.2')

# ─────────────────────────────────────────────────────────────────────
# 6. НОВИЙ ПІДРОЗДІЛ 1.3 — ВЕБ-ІНТЕРФЕЙС СИСТЕМИ
# ─────────────────────────────────────────────────────────────────────

# Вставляємо після розширення 1.2 — знаходимо "2. РЕАЛІЗАЦІЯ ПЛАНУВАЛЬНИКА"
planner_para = find_para(doc, '2. РЕАЛІЗАЦІЯ ПЛАНУВАЛЬНИКА')
if planner_para:
    web_section = [
        add_para('1.3. Опис веб-інтерфейсу системи', bold=True),
        add_para(),
        add_para(
            'Система надає веб-інтерфейс на базі HTML/CSS/JavaScript без '
            'окремого frontend-фреймворку. Усі сторінки повертаються '
            'сервером FastAPI безпосередньо як HTMLResponse. '
            'JavaScript-код взаємодіє з REST API через стандартний Fetch API '
            'та EventSource для SSE-потоку логів.'
        ),
        add_para(),
        add_para(
            'Головна сторінка (GET /) надає студенту повний набір '
            'інструментів для подачі задачі. Центральне місце займає Monaco '
            'Editor — той самий редактор коду, що вбудований у Visual Studio '
            'Code, з підсвіткою синтаксису Python. Поруч розміщено форму з '
            'параметрами: вибір runtime-профілю зі списку, поля CPU та RAM, '
            'опція завантаження .py / .zip файлів або URL git-репозиторію. '
            'Результати виконання (статус та логи) відображаються на тій '
            'самій сторінці без перезавантаження (рис. 3.6).'
        ),
        add_para(),
        add_placeholder(
            'Головна сторінка: Monaco Editor з кодом Python зліва, '
            'форма вибору runtime / CPU / RAM / файл / git справа, '
            'секція статусу знизу'
        ),
        add_para('Рисунок 3.6 — Головна сторінка веб-інтерфейсу',
                 italic=True, center=True),
        add_para(),
        add_para(
            'Сторінка статусу задачі відображає поточний стан у реальному '
            'часі. При статусі queued показано позицію в черзі. Після '
            'переходу до статусу running JavaScript відкриває SSE-з\'єднання '
            'до GET /jobs/{id}/logs/stream і виводить кожен рядок логу '
            'миттєво у текстову область. Після отримання SSE-події done '
            'з\'єднання закривається і відображається список '
            'артефактів-файлів для завантаження (рис. 3.7).'
        ),
        add_para(),
        add_placeholder(
            'Сторінка статусу: поле статусу (running), '
            'прокручуваний SSE-лог з виводом скрипту в реальному часі, '
            'кнопка "Скасувати"'
        ),
        add_para('Рисунок 3.7 — Сторінка статусу задачі з SSE-стрімінгом логів',
                 italic=True, center=True),
        add_para(),
        add_para(
            'Адміністративна панель (GET /admin) захищена перевіркою '
            'API-ключа при завантаженні сторінки (GET /admin/check). '
            'Вона надає викладачеві три інструменти: '
            '1) таблицю задач із фільтрацією за student_id та статусом і '
            'кнопками примусового зупинення; '
            '2) форму генерації нових API-ключів студентів; '
            '3) таблицю існуючих ключів із можливістю видалення (рис. 3.8).'
        ),
        add_para(),
        add_placeholder(
            'Адміністративна панель: таблиця задач з колонками '
            'job_id / student_id / status / runtime / GPU / час, '
            'форма "Новий ключ" з полем student_id'
        ),
        add_para('Рисунок 3.8 — Адміністративна панель системи',
                 italic=True, center=True),
        add_para(),
    ]
    # Вставляємо ПЕРЕД розділом 2 (тобто після останнього абзацу 1.2)
    chain_insert_after(planner_para._p.getprevious(), web_section)
    print('  [NEW] додано підрозділ 1.3 «Веб-інтерфейс»')
else:
    print('  [WARN] не знайдено "2. РЕАЛІЗАЦІЯ ПЛАНУВАЛЬНИКА"')

# ─────────────────────────────────────────────────────────────────────
# 7. РОЗШИРЕННЯ ТИЖНЯ 4 — ДОКУМЕНТАЦІЯ
# ─────────────────────────────────────────────────────────────────────

deploy_para = find_para(doc, 'DEPLOY.md — покрокова інструкція')
if deploy_para:
    week4_ext = [
        add_para(),
        add_para(
            'Шість кроків DEPLOY.md охоплюють: '
            '1) встановлення NVIDIA-драйверів та NVIDIA Container Toolkit; '
            '2) встановлення Docker і Docker Compose; '
            '3) клонування репозиторію і налаштування .env; '
            '4) запуск redis і master через docker compose up; '
            '5) запуск worker з параметром GPU_COUNT; '
            '6) перевірку системи через GET /health. '
            'Для кожного кроку наведено дослівну команду і пояснення.'
        ),
        add_para(),
        add_para(
            'benchmarks/ містить вісім відтворюваних Python-скриптів: '
            'lifecycle.py (вимірює повний цикл від submit до завершення), '
            'concurrency.py (паралельне навантаження 1/5/10/20 задач), '
            'sse_vs_polling.py (порівняння latency), '
            'runtime_compare.py (matmul 4000×4000 по профілях), '
            'cold_warm.py (час холодного та теплого старту контейнера), '
            'security_audit.py (23 автоматизованих атаки), '
            'diagram.py (генерація архітектурних діаграм). '
            'Кожен скрипт генерує CSV у benchmarks/results/ і PNG у '
            'benchmarks/results/plots/ для вставки в документацію.'
        ),
        add_para(),
        add_para(
            'Окремий бенчмарк cold_warm.py порівнює час запуску першого '
            '(cold) та наступних (warm) контейнерів для кожного '
            'runtime-профілю. Cold start включає розпакування шарів образу '
            'Docker; warm start використовує кешовані шари. '
            'Для pytorch-cu121 mid-term warm start на 40–60% швидший за '
            'cold start завдяки кешуванню Docker-шарів на диску хоста.'
        ),
        add_para(),
    ]
    chain_insert_after(deploy_para._p, week4_ext)
    print('  [EXP] розширено розділ «Документація» Тижня 4')

# ─────────────────────────────────────────────────────────────────────
# 8. ПОВНИЙ ТЕКСТ ВИСНОВКІВ
# ─────────────────────────────────────────────────────────────────────

vysnovky = find_para(doc, 'ВИСНОВКИ')
if vysnovky:
    conclusions = [
        add_para(),
        add_para(
            'У рамках переддипломної практики розроблено повнофункціональний '
            'сервіс дистанційного виконання обчислювальних задач на GPU в '
            'ізольованих контейнерах. Усі заплановані технічні завдання '
            'виконано у повному обсязі відповідно до індивідуального '
            'завдання.'
        ),
        add_para(),
        add_para(
            'На першому тижні проведено глибокий аналіз предметної галузі: '
            'досліджено чотири хмарні GPU-платформи (Google Colab, Kaggle, '
            'RunPod, CoreWeave), порівняно підходи до ізоляції середовищ '
            '(VM, Docker, LXC) та три системи планування задач (Slurm, '
            'Kubernetes, Nomad). Вивчено архітектуру NVIDIA Container '
            'Toolkit і механізм конфліктів версій CUDA/cuDNN. Встановлено, '
            'що жодна з готових систем не є оптимальною для одного '
            'кафедрального сервера без виділеної DevOps-команди; '
            'оптимальним є власний планувальник на базі FastAPI + Redis + '
            'Docker.'
        ),
        add_para(),
        add_para(
            'На другому тижні спроектовано мікросервісну архітектуру з трьох '
            'компонентів: Master Server (FastAPI), Worker та Redis. '
            'Обґрунтовано вибір технологічного стеку: FastAPI з asyncio для '
            'SSE-стрімінгу без блокування, Redis з атомарними операціями '
            'WATCH/MULTI/EXEC для виключення race condition при резервуванні '
            'GPU, Docker з прапором --gpus device=N для ексклюзивного '
            'доступу до конкретної карти. Розроблено семирівневу модель '
            'безпеки за принципом defense in depth: AST-аналіз, мережева '
            'ізоляція, Linux capabilities, cgroups v2, zip-slip перевірка, '
            'ліміт черги та API-ключі.'
        ),
        add_para(),
        add_para(
            'На третьому тижні реалізовано повнофункціональну систему '
            'обсягом 1308 рядків виробничого коду. Введено три '
            'runtime-профілі (pytorch-cu121, pytorch-cu118, tensorflow) '
            'для вирішення проблеми конфліктів версій CUDA. '
            'Реалізовано SSE-стрімінг логів у реальному часі та '
            'веб-інтерфейс з Monaco Editor і адміністративною панеллю. '
            'Проведено комплексне тестування: 62 автоматизованих тести у '
            '13 класах pytest, вісім бенчмарків, аудит безпеки з 23 '
            'перевірками — усі успішно пройдено, нуль race conditions '
            'при 20 паралельних задачах.'
        ),
        add_para(),
        add_para(
            'На четвертому тижні підготовлено повний пакет технічної '
            'документації: README.md, DEPLOY.md із шестикроковою '
            'інструкцією розгортання, .env.example з коментованими '
            'змінними, вісім відтворюваних бенчмаркових скриптів. '
            'Оформлено матеріали до пояснювальної записки дипломної роботи.'
        ),
        add_para(),
        add_para('Кількісні результати виконаної роботи:', bold=True),
        add_para('— 1308 рядків виробничого коду (master.py + worker.py);'),
        add_para('— 21 REST API ендпоінт;'),
        add_para('— 62 автоматизованих тести (pytest) у 13 класах;'),
        add_para('— 8 бенчмарків та 8 PNG-графіків результатів;'),
        add_para('— 7 незалежних рубежів безпеки;'),
        add_para('— 3 runtime-профілі;'),
        add_para('— 38 git-комітів.'),
        add_para(),
        add_para(
            'Результати бенчмарків підтвердили ефективність обраних рішень. '
            'SSE-стрімінг у 4.4 рази швидший за polling з інтервалом '
            '5 секунд і генерує значно менший мережевий трафік. '
            'Runtime-профіль pytorch-cu121 найшвидший при матричному '
            'множенні 4000×4000 (49.1 мс проти 73.1 мс для pytorch-cu118 '
            'і 163.6 мс для tensorflow). Час очікування задачі в черзі '
            'менше 1% від загального часу виконання — системний overhead '
            'незначний. Cold start контейнера на 40–60% довший за warm '
            'start завдяки кешуванню Docker-шарів.'
        ),
        add_para(),
        add_para(
            'Розроблена система повністю вирішує поставлену задачу: '
            'студенти отримали зручний доступ до GPU через веб-інтерфейс '
            'або REST API без ручного налаштування середовища і без '
            'конфліктів між задачами. Кожне завдання виконується в '
            'ізольованому Docker-контейнері з обраним CUDA-стеком, GPU '
            'розподіляються чесно через Redis-чергу, а довільний '
            'студентський код захищений сімома незалежними рубежами.'
        ),
        add_para(),
        add_para(
            'Перспективні напрями розвитку: підтримка NVIDIA MIG для '
            'поділу одного GPU між кількома задачами; інтеграція з LMS '
            'кафедри для автоматичного прийняття лабораторних робіт; '
            'розширення каталогу runtime-профілів (JAX, RAPIDS cuDF, '
            'torch.compile); пріоритетна черга для диференціації задач '
            'студентів і викладачів; Prometheus-метрики для моніторингу '
            'навантаження в реальному часі.'
        ),
        add_para(),
    ]
    chain_insert_after(vysnovky._p, conclusions)
    print('  [NEW] написано ВИСНОВКИ')

# ─────────────────────────────────────────────────────────────────────
# 9. ДОДАТОК А — ДОКУМЕНТАЦІЯ REST API
# ─────────────────────────────────────────────────────────────────────

appendix = find_para(doc, 'Додаток А')
if appendix:
    api_doc = [
        add_para(),
        add_para('ПОВНА ДОКУМЕНТАЦІЯ REST API (21 ЕНДПОІНТ)', bold=True),
        add_para(),
        add_para(
            'У цьому додатку наведено повну документацію всіх REST API '
            'ендпоінтів Master Server. Для кожного ендпоінта вказано метод, '
            'шлях, опис, параметри та приклад відповіді. Автоматична '
            'OpenAPI-документація доступна за адресою /docs (Swagger UI) '
            'та /redoc (ReDoc).'
        ),
        add_para(),

        # ── Загальні ──
        add_para('А.1. Публічні ендпоінти моніторингу', bold=True),
        add_para(),

        add_para('GET /health', bold=True),
        add_para(
            'Перевірка стану сервісу. Повертає статус Master Server, '
            'доступність Redis і кількість знайдених GPU.'
        ),
        add_para('Відповідь 200: {"status": "ok", "redis": "ok", "gpus": 2}'),
        add_para(),

        add_para('GET /runtimes', bold=True),
        add_para(
            'Перелік доступних runtime-профілів з назвами Docker-образів. '
            'Використовується веб-інтерфейсом для побудови списку вибору.'
        ),
        add_para(
            'Відповідь 200: {"runtimes": {"pytorch-cu121": '
            '"pytorch/pytorch:2.2.2-cuda12.1-cudnn8-runtime", '
            '"pytorch-cu118": "pytorch/pytorch:2.1.2-cuda11.8-cudnn8-runtime", '
            '"tensorflow": "tensorflow/tensorflow:2.15.0-gpu"}}'
        ),
        add_para(),

        add_para('GET /gpus', bold=True),
        add_para(
            'Поточний стан кожної GPU: статус (idle/busy), '
            'job_id поточної задачі або null.'
        ),
        add_para(
            'Відповідь 200: {"gpus": [{"gpu_id": 0, "status": "busy", '
            '"job_id": "abc123"}, {"gpu_id": 1, "status": "idle", '
            '"job_id": null}]}'
        ),
        add_para(),

        add_para('GET /cluster-status', bold=True),
        add_para(
            'Агрегована статистика: кількість GPU, поточне навантаження, '
            'довжина черги, лічильники задач за статусами.'
        ),
        add_para(
            'Відповідь 200: {"total_gpus": 2, "idle_gpus": 1, '
            '"queue_length": 3, '
            '"jobs": {"queued": 3, "running": 1, "completed": 42, '
            '"failed": 2, "cancelled": 1}}'
        ),
        add_para(),

        # ── Подача ──
        add_para('А.2. Подача задач', bold=True),
        add_para(),

        add_para('POST /submit-job (Content-Type: application/json)', bold=True),
        add_para(
            'Подача задачі у форматі JSON. Заголовок X-API-Key '
            'необов\'язковий: якщо відсутній, student_id береться з тіла '
            'запиту; якщо присутній, student_id береться з ключа незалежно '
            'від тіла (захист від підробки).'
        ),
        add_para('Параметри тіла (всі, крім code, мають значення за замовчуванням):'),
        add_para('  code (str, обов\'язковий) — Python-код. Ліміт 200 КБ.'),
        add_para('  student_id (str) — ідентифікатор студента.'),
        add_para(
            '  runtime (str) — профіль: pytorch-cu121 | pytorch-cu118 | '
            'tensorflow. За замовч. pytorch-cu121.'
        ),
        add_para('  cpus (float) — кількість CPU, 0.1–8.0. За замовч. 2.0.'),
        add_para(
            '  memory (str) — RAM у форматі Docker (512m, 2g тощо). '
            'За замовч. 2g.'
        ),
        add_para(
            'Відповідь 201: {"job_id": "abc123", "status": "queued", '
            '"runtime": "pytorch-cu121", "position": 2}'
        ),
        add_para(
            'Помилки: 400 — небезпечний код (AST-аналіз), '
            '400 — невідомий runtime, 429 — перевищено '
            'MAX_QUEUED_PER_STUDENT.'
        ),
        add_para(),

        add_para('POST /submit-job-form (multipart/form-data)', bold=True),
        add_para(
            'Подача через форму. Підтримує три режими: код як текст, '
            '.py файл, git-репозиторій. Режими мають пріоритет: '
            'git_repo > file > code.'
        ),
        add_para('Поля форми:'),
        add_para('  code (str) — Python-код (використовується, якщо файл не завантажено).'),
        add_para('  file (UploadFile) — .py файл (пріоритет над code).'),
        add_para('  git_repo (str) — URL публічного git-репозиторію.'),
        add_para('  entrypoint (str) — стартовий файл при git-режимі. За замовч. main.py.'),
        add_para('  dataset (UploadFile) — .zip файл з даними (→ /workspace/data/).'),
        add_para('  requirements (UploadFile) — requirements.txt для pip install.'),
        add_para('  runtime, cpus, memory, student_id — аналогічно /submit-job.'),
        add_para('Відповідь 201: аналогічно /submit-job.'),
        add_para(),

        # ── Задачі ──
        add_para('А.3. Управління задачами', bold=True),
        add_para(),

        add_para('GET /jobs', bold=True),
        add_para(
            'Перелік задач з опціональним фільтруванням. '
            'Query-параметри: student_id (str), status (str), '
            'limit (int, за замовч. 50).'
        ),
        add_para('Відповідь 200: {"jobs": [{...}, ...]}'),
        add_para(),

        add_para('GET /jobs/{job_id}', bold=True),
        add_para(
            'Деталі задачі: статус, параметри, gpu_id, час постановки, '
            'старту, завершення, позиція в черзі (для queued).'
        ),
        add_para(
            'Відповідь 200: {"job_id": "abc123", "status": "running", '
            '"student_id": "st01", "runtime": "pytorch-cu121", '
            '"cpus": 2.0, "memory": "2g", "gpu_id": 0, '
            '"queued_at": "2026-05-18T10:00:00", '
            '"started_at": "2026-05-18T10:00:01", '
            '"finished_at": null, "position": null}'
        ),
        add_para('Помилка: 404 — задача не знайдена.'),
        add_para(),

        add_para('POST /jobs/{job_id}/cancel', bold=True),
        add_para(
            'Скасування задачі. Для queued — видалення з черги. '
            'Для running — docker stop контейнера.'
        ),
        add_para(
            'Відповідь 200: {"status": "cancelled"}. '
            'Помилки: 404, 400 — задача вже завершена.'
        ),
        add_para(),

        add_para('GET /jobs/{job_id}/logs/stream', bold=True),
        add_para(
            'SSE-стрімінг логів виконання. '
            'Content-Type: text/event-stream. '
            'Повертає рядки stdout/stderr у форматі SSE. '
            'Остання подія: data: done — сигналізує про завершення.'
        ),
        add_para('Приклад відповіді:'),
        add_para('  data: stdout:Epoch 1/10, loss=0.4231'),
        add_para('  data: stdout:Epoch 2/10, loss=0.3812'),
        add_para('  data: done'),
        add_para(),

        add_para('GET /jobs/{job_id}/artifacts', bold=True),
        add_para(
            'Перелік файлів у /workspace/output після завершення задачі.'
        ),
        add_para(
            'Відповідь 200: {"job_id": "abc123", '
            '"artifacts": ["model.pt", "loss_curve.png", "metrics.csv"]}'
        ),
        add_para(),

        add_para('GET /jobs/{job_id}/artifacts/{filename}', bold=True),
        add_para(
            'Завантаження конкретного файлу-артефакту. '
            'Захищений від path traversal через перевірку .resolve().'
        ),
        add_para(
            'Відповідь 200: FileResponse. '
            'Помилки: 404 — файл не знайдено, 400 — path traversal.'
        ),
        add_para(),

        add_para('DELETE /jobs/cleanup', bold=True),
        add_para(
            'Видалення даних завершених задач із Redis. '
            'Query-параметр: older_than_hours (int, за замовч. 24).'
        ),
        add_para('Відповідь 200: {"deleted": 15}'),
        add_para(),

        # ── Адмін ──
        add_para('А.4. Адміністративні ендпоінти', bold=True),
        add_para(
            'Усі ендпоінти /admin/* вимагають заголовок '
            'X-API-Key зі значенням ADMIN_API_KEY. '
            'Відповідь 401 — ключ відсутній або невірний.'
        ),
        add_para(),

        add_para('GET /admin/check', bold=True),
        add_para(
            'Перевірка адміністративного ключа. '
            'Використовується веб-панеллю при завантаженні сторінки.'
        ),
        add_para(
            'Відповідь 200: {"status": "authorized"}. '
            'Відповідь 401: {"detail": "unauthorized"}.'
        ),
        add_para(),

        add_para('GET /admin/stats', bold=True),
        add_para(
            'Агрегована статистика: кількість задач за student_id, '
            'використання runtime-профілів, погодинні лічильники.'
        ),
        add_para(
            'Відповідь 200: {"by_student": {"st01": 12, "st02": 8}, '
            '"by_runtime": {"pytorch-cu121": 15, "tensorflow": 5}, '
            '"by_hour": {"2026-05-18T10": 3, ...}}'
        ),
        add_para(),

        add_para('POST /admin/jobs/{job_id}/kill', bold=True),
        add_para(
            'Примусове зупинення будь-якої задачі незалежно від статусу. '
            'Надсилає docker kill до контейнера.'
        ),
        add_para(
            'Відповідь 200: {"status": "killed"}. Помилка: 404.'
        ),
        add_para(),

        add_para('POST /admin/keys', bold=True),
        add_para('Генерація нового API-ключа для студента.'),
        add_para('Тіло (JSON): {"student_id": "st01", "expires_days": 90}'),
        add_para(
            'Відповідь 201: {"api_key": "sk-a1b2c3...", '
            '"student_id": "st01", "expires_at": "2026-08-16T00:00:00"}'
        ),
        add_para(),

        add_para('GET /admin/keys', bold=True),
        add_para('Перелік усіх API-ключів з прив\'язаними student_id та строками дії.'),
        add_para(
            'Відповідь 200: {"keys": [{"api_key": "sk-***c3d4", '
            '"student_id": "st01", "expires_at": "2026-08-16"}, ...]}'
        ),
        add_para(),

        add_para('DELETE /admin/keys/{api_key}', bold=True),
        add_para('Анулювання API-ключа студента.'),
        add_para(
            'Відповідь 200: {"deleted": true}. Помилка: 404.'
        ),
        add_para(),

        # ── UI ──
        add_para('А.5. Ендпоінти веб-інтерфейсу', bold=True),
        add_para(),

        add_para('GET / (HTMLResponse)', bold=True),
        add_para(
            'Головна сторінка веб-інтерфейсу з Monaco Editor та формою '
            'подачі задачі. Не потребує автентифікації.'
        ),
        add_para(),

        add_para('GET /admin (HTMLResponse)', bold=True),
        add_para(
            'Адміністративна панель. Захист реалізовано на рівні JavaScript '
            '(перевірка X-API-Key через GET /admin/check при завантаженні); '
            'сам HTML-файл доступний без ключа.'
        ),
        add_para(),
    ]
    chain_insert_after(appendix._p, api_doc)
    print('  [NEW] заповнено Додаток А')

# ─────────────────────────────────────────────────────────────────────
# ЗБЕРЕЖЕННЯ
# ─────────────────────────────────────────────────────────────────────

doc.save(DOCX_OUT)
print(f'\nГотово! Збережено: {DOCX_OUT}')
